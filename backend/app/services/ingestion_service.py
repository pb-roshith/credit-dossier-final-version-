"""
Ingestion Service — processes uploaded files using Mistral's OCR API
for high-accuracy text extraction (tables, scanned docs, complex layouts).

Documents are stored at the deal level and can be linked to multiple sections.
SHA-256 hashing prevents re-processing duplicate files.

Fallback: local extraction via pypdf/python-docx if Mistral OCR fails.
"""

from __future__ import annotations

import hashlib
import io
import logging
import os
import uuid
from pathlib import Path
from typing import Optional

import httpx
from sqlalchemy.orm import Session

from app.config import settings
from app.models.deal import Deal
from app.models.document import DealDocument, SectionDocumentLink
from app.models.upload import Upload

logger = logging.getLogger(__name__)

# Lazy Mistral client
_mistral_client = None


def _get_mistral():
    global _mistral_client
    if _mistral_client is None:
        from mistralai.client import Mistral
        _mistral_client = Mistral(api_key=settings.MISTRAL_API_KEY)
    return _mistral_client


# ── Mistral OCR Extraction ──────────────────────────────────────


async def extract_text_with_mistral_ocr(
    file_bytes: bytes, filename: str
) -> tuple[str, int]:
    """
    Use Mistral's OCR API for high-accuracy text extraction.
    Produces clean Markdown with properly formatted tables.

    Returns (extracted_markdown, page_count).

    Flow:
    1. Upload file to Mistral Files API (purpose="ocr")
    2. Call ocr.process_async() with mistral-ocr-latest
    3. Concatenate all page.markdown outputs
    4. Delete the temporary file from Mistral storage
    """
    client = _get_mistral()

    # 1. Upload to Mistral file storage
    uploaded = await client.files.upload_async(
        file={"file_name": filename, "content": file_bytes},
        purpose="ocr",
    )

    logger.info(f"Uploaded {filename} to Mistral files (id={uploaded.id})")

    # 2. Process with OCR
    ocr_response = await client.ocr.process_async(
        model="mistral-ocr-latest",
        document={"type": "file_id", "file_id": uploaded.id},
    )

    # 3. Concatenate all pages
    pages_md = []
    for i, page in enumerate(ocr_response.pages):
        pages_md.append(f"<!-- Page {i + 1} -->\n{page.markdown}")

    result = "\n\n".join(pages_md)
    page_count = len(ocr_response.pages)

    logger.info(
        f"OCR complete for {filename}: {page_count} pages, "
        f"{len(result)} chars extracted"
    )

    # 4. Cleanup: delete the temp file from Mistral storage
    try:
        await client.files.delete_async(file_id=uploaded.id)
    except Exception as e:
        logger.warning(f"Could not delete temp Mistral file {uploaded.id}: {e}")

    return result, page_count


# ── Local Fallback Extraction ───────────────────────────────────


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Fallback PDF extraction using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Fallback DOCX extraction using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)
    return "\n".join(paragraphs)


def _local_fallback_extract(file_bytes: bytes, filename: str) -> str:
    """Fallback: use local libraries if Mistral OCR fails."""
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_text_from_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            return _extract_text_from_docx(file_bytes)
        else:
            return f"[Could not extract text from {filename}]"
    except Exception as e:
        logger.error(f"Local fallback extraction failed for {filename}: {e}")
        return f"[Extraction failed for {filename}]"


# Keep this sync version for template extraction (non-OCR)
def extract_text_preview(file_bytes: bytes, filename: str) -> str:
    """Extract text preview for templates. Uses local extraction only."""
    ext = Path(filename).suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_text_from_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            return _extract_text_from_docx(file_bytes)
        elif ext in (".txt", ".md", ".csv", ".json"):
            return file_bytes.decode("utf-8", errors="replace")
        else:
            return f"[File: {filename}]"
    except Exception:
        return f"[Could not preview {filename}]"


# ── Deal Document Processing (New Architecture) ────────────────


# Extensions that should be processed via Mistral OCR
OCR_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".xlsx", ".xls",
    ".pptx", ".png", ".jpg", ".jpeg", ".tiff", ".bmp",
}

# Extensions that are plain text (no OCR needed)
PLAIN_TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json"}


class IngestionService:
    """
    Handles document upload, OCR processing, and section linking.

    Documents are stored at the deal level (DealDocument).
    Sections reference documents via SectionDocumentLink (many-to-many).
    SHA-256 hashing prevents re-processing duplicate files.
    """

    @staticmethod
    async def _save_file(
        file_bytes: bytes, filename: str, deal_id: str
    ) -> str:
        """Save uploaded file to disk. Returns the file path."""
        upload_dir = settings.upload_path / deal_id
        upload_dir.mkdir(parents=True, exist_ok=True)

        safe_name = f"{uuid.uuid4().hex[:8]}_{filename}"
        file_path = upload_dir / safe_name

        with open(file_path, "wb") as f:
            f.write(file_bytes)

        return str(file_path)

    @staticmethod
    async def process_deal_document(
        db: Session,
        deal_id: str,
        file_bytes: bytes,
        filename: str,
        source_type: str = "file",
        note: str | None = None,
        url: str | None = None,
        text_content: str | None = None,
    ) -> DealDocument:
        """
        Process and store a document at deal level.

        1. Compute SHA-256 hash → check for duplicate in this deal
        2. Save file to disk
        3. Extract text via Mistral OCR (or plain read for text files)
        4. Create DealDocument record

        Returns existing document if hash matches (no re-processing).
        """
        # 1. Compute file hash for dedup
        file_hash = hashlib.sha256(file_bytes).hexdigest()

        existing = (
            db.query(DealDocument)
            .filter(
                DealDocument.deal_id == deal_id,
                DealDocument.file_hash == file_hash,
            )
            .first()
        )
        if existing:
            logger.info(
                f"Duplicate detected for {filename} in deal {deal_id} "
                f"(hash={file_hash[:12]}…). Reusing doc {existing.id}"
            )
            return existing

        # 2. Save file to disk
        file_path = await IngestionService._save_file(
            file_bytes, filename, deal_id
        )

        # 3. Extract text
        ext = Path(filename).suffix.lower()
        extracted_text = ""
        extraction_method = "unknown"
        page_count = 0

        if ext in PLAIN_TEXT_EXTENSIONS:
            extracted_text = file_bytes.decode("utf-8", errors="replace")
            extraction_method = "plain_text"
            page_count = 1
        elif ext in OCR_EXTENSIONS:
            try:
                extracted_text, page_count = (
                    await extract_text_with_mistral_ocr(file_bytes, filename)
                )
                extraction_method = "mistral_ocr"
            except Exception as e:
                logger.error(
                    f"Mistral OCR failed for {filename}, "
                    f"falling back to local: {e}"
                )
                extracted_text = _local_fallback_extract(file_bytes, filename)
                extraction_method = "local_fallback"
                page_count = 0
        else:
            extracted_text = f"[Unsupported file type: {ext}]"
            extraction_method = "unsupported"

        # 4. Create DealDocument record
        doc = DealDocument(
            deal_id=deal_id,
            source_type=source_type,
            filename=filename,
            file_path=file_path,
            file_hash=file_hash,
            url=url,
            text_content=text_content,
            note=note,
            extracted_text=extracted_text[:200_000] if extracted_text else None,
            extraction_method=extraction_method,
            page_count=page_count,
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        logger.info(
            f"Created DealDocument {doc.id} for {filename} "
            f"(method={extraction_method}, pages={page_count}, "
            f"chars={len(extracted_text)})"
        )
        return doc

    @staticmethod
    async def process_url_document(
        db: Session,
        deal_id: str,
        url: str,
        note: str | None = None,
    ) -> DealDocument:
        """Fetch content from a URL and store as a deal document."""
        fetched_bytes = b""
        filename = url.split("/")[-1] or "downloaded_file.txt"

        try:
            async with httpx.AsyncClient(timeout=30.0) as http_client:
                response = await http_client.get(url)
                response.raise_for_status()
                content_type = response.headers.get("content-type", "")
                if "text" in content_type or "json" in content_type:
                    fetched_bytes = response.text.encode("utf-8")
                    if not filename.endswith((".txt", ".json", ".csv", ".md")):
                        filename = filename + ".txt"
                else:
                    fetched_bytes = response.content
        except Exception as e:
            logger.error(f"Failed to fetch URL {url}: {e}")
            # Store a stub document
            doc = DealDocument(
                deal_id=deal_id,
                source_type="url",
                url=url,
                note=note,
                extracted_text=f"[Failed to fetch URL: {e}]",
                extraction_method="error",
            )
            db.add(doc)
            db.commit()
            db.refresh(doc)
            return doc

        return await IngestionService.process_deal_document(
            db=db,
            deal_id=deal_id,
            file_bytes=fetched_bytes,
            filename=filename,
            source_type="url",
            note=note,
            url=url,
        )

    @staticmethod
    async def process_text_document(
        db: Session,
        deal_id: str,
        text_content: str,
        note: str | None = None,
    ) -> DealDocument:
        """Store pasted text as a deal document. No OCR needed."""
        text_bytes = text_content.encode("utf-8")
        file_hash = hashlib.sha256(text_bytes).hexdigest()

        # Check for duplicate
        existing = (
            db.query(DealDocument)
            .filter(
                DealDocument.deal_id == deal_id,
                DealDocument.file_hash == file_hash,
            )
            .first()
        )
        if existing:
            return existing

        doc = DealDocument(
            deal_id=deal_id,
            source_type="text",
            text_content=text_content,
            note=note,
            file_hash=file_hash,
            extracted_text=text_content[:200_000],
            extraction_method="plain_text",
            page_count=1,
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        return doc

    # ── Section Document Linking ─────────────────────────────────

    @staticmethod
    def link_documents_to_section(
        db: Session,
        section_id: str,
        document_ids: list[str],
    ) -> list[SectionDocumentLink]:
        """Link one or more deal documents to a section (idempotent)."""
        links = []
        for doc_id in document_ids:
            # Check if already linked
            existing = (
                db.query(SectionDocumentLink)
                .filter(
                    SectionDocumentLink.section_id == section_id,
                    SectionDocumentLink.document_id == doc_id,
                )
                .first()
            )
            if existing:
                links.append(existing)
                continue

            link = SectionDocumentLink(
                section_id=section_id,
                document_id=doc_id,
            )
            db.add(link)
            links.append(link)

        db.commit()
        for link in links:
            db.refresh(link)

        return links

    @staticmethod
    def unlink_document_from_section(
        db: Session,
        section_id: str,
        document_id: str,
    ) -> bool:
        """Remove a document link from a section."""
        link = (
            db.query(SectionDocumentLink)
            .filter(
                SectionDocumentLink.section_id == section_id,
                SectionDocumentLink.document_id == document_id,
            )
            .first()
        )
        if not link:
            return False

        db.delete(link)
        db.commit()
        return True

    @staticmethod
    def delete_deal_document(db: Session, document_id: str) -> bool:
        """Delete a deal document, its file, and all section links."""
        doc = (
            db.query(DealDocument)
            .filter(DealDocument.id == document_id)
            .first()
        )
        if not doc:
            return False

        # Remove file from disk
        if doc.file_path and os.path.exists(doc.file_path):
            try:
                os.remove(doc.file_path)
            except Exception as e:
                logger.error(f"Failed to delete file {doc.file_path}: {e}")

        # Cascade deletes section_links automatically
        db.delete(doc)
        db.commit()
        return True

    # ── Legacy Upload Support ────────────────────────────────────

    @staticmethod
    async def save_uploaded_file(
        file_bytes: bytes, filename: str, deal_id: str
    ) -> str:
        """Legacy: Save uploaded file to disk."""
        return await IngestionService._save_file(file_bytes, filename, deal_id)

    @staticmethod
    async def process_file_upload(
        db: Session,
        section_id: str,
        deal_id: str,
        file_bytes: bytes,
        filename: str,
        note: str | None = None,
    ) -> Upload:
        """Legacy: Process a file upload at section level."""
        file_path = await IngestionService._save_file(
            file_bytes, filename, deal_id
        )

        # Use local extraction for legacy uploads
        extracted = extract_text_preview(file_bytes, filename)

        upload = Upload(
            section_id=section_id,
            source_type="file",
            filename=filename,
            file_path=file_path,
            note=note,
            extracted_text=extracted[:50000] if extracted else None,
        )
        db.add(upload)
        db.commit()
        db.refresh(upload)
        return upload

    @staticmethod
    async def process_url_upload(
        db: Session,
        section_id: str,
        deal_id: str,
        url: str,
        note: str | None = None,
    ) -> Upload:
        """Legacy: Process a URL upload at section level."""
        extracted = ""
        try:
            async with httpx.AsyncClient(timeout=30.0) as http_client:
                response = await http_client.get(url)
                response.raise_for_status()
                content_type = response.headers.get("content-type", "")
                if "text" in content_type or "json" in content_type:
                    extracted = response.text
                else:
                    fetched_bytes = response.content
                    filename_from_url = url.split("/")[-1] or "downloaded_file"
                    extracted = extract_text_preview(
                        fetched_bytes, filename_from_url
                    )
        except Exception as e:
            logger.error(f"Failed to fetch URL {url}: {e}")
            extracted = f"[Failed to fetch URL: {e}]"

        upload = Upload(
            section_id=section_id,
            source_type="url",
            url=url,
            note=note,
            extracted_text=extracted[:50000] if extracted else None,
        )
        db.add(upload)
        db.commit()
        db.refresh(upload)
        return upload

    @staticmethod
    async def process_text_upload(
        db: Session,
        section_id: str,
        deal_id: str,
        text_content: str,
        note: str | None = None,
    ) -> Upload:
        """Legacy: Process a text upload at section level."""
        upload = Upload(
            section_id=section_id,
            source_type="text",
            text_content=text_content,
            note=note,
            extracted_text=text_content[:50000],
        )
        db.add(upload)
        db.commit()
        db.refresh(upload)
        return upload

    @staticmethod
    def delete_upload(db: Session, upload_id: str) -> bool:
        """Legacy: Delete an upload."""
        upload = db.query(Upload).filter(Upload.id == upload_id).first()
        if not upload:
            return False

        if upload.file_path and os.path.exists(upload.file_path):
            try:
                os.remove(upload.file_path)
            except Exception as e:
                logger.error(f"Failed to delete file {upload.file_path}: {e}")

        db.delete(upload)
        db.commit()
        return True
